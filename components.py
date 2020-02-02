import requests
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
from os import path
from styles import MONTSERRAT_REGULAR, CAUTION_STYLE, CAUTION_CODE, \
    HEADING_1, HEADING_2, HEADING_3, NOTE_STYLE, NOTE_CODE, CODE_BLUE, CODE_PINK, CODE_MAROON, CODE_LIGHT_BLUE, \
    CODE_DARK_GREEN, CODE_PURPLE


def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)


def unordered_list(document, el):
    for c in el.findAll('li'):
        p = paragraph(document, c, True)
        p.style = 'List Bullet'


def ordered_list(document, el):
    for c in el.findAll('li'):
        p = paragraph(document, c, True)
        p.style = 'List Number'


def paragraph(document, el, indent=False):
    p = document.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(40)
    p.style = MONTSERRAT_REGULAR
    for c in el.contents:
        if isinstance(c, str):
            p.add_run(c.replace('\n', ' '))
        if c.name == 'a':
            p.add_run(c)
        elif c.name == 'b' or c.name == 'strong':
            bold(p, c.text)
        elif c.name == 'i' or c.name == 'em':
            italic(p, c.text)
        elif c.name == 'code':
            code(p, c.text)
        elif c.name == 'ul':
            unordered_list(document, c)
            p = document.add_paragraph()
            p.style = MONTSERRAT_REGULAR
            indent_size = Pt(40)
            if indent:
                p.paragraph_format.left_indent = indent_size
        elif c.name == 'ol':
            ordered_list(document, c)
            p = document.add_paragraph()
            p.style = MONTSERRAT_REGULAR
            indent_size = Pt(40)
            if indent:
                p.paragraph_format.left_indent = indent_size
        elif c.name == 'p':
            if c.has_attr('class'):
                if c['class'][0] == 'note':
                    note(document, c)
                elif c['class'][0] == 'caution':
                    caution(document, c)
            else:
                p = paragraph(document, c, indent)
        elif el.name == 'img':
            img(document, el.src)
    return p


def data_dictionary(document, el):
    for e in el.contents:
        if e.name == 'dt':
            paragraph(document, e)
        elif e.name == 'dd':
            paragraph(document, e, True)


def note(document, el):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cells([cell], "#E1F5FE")
    p = cell.paragraphs[0]
    p.style = NOTE_STYLE
    pf = p.paragraph_format
    pf.space_before = Pt(15)
    pf.space_after = Pt(15)
    pf.left_indent = Pt(20)
    pf.right_indent = Pt(15)

    for c in el.contents:
        if isinstance(c, str):
            p.add_run(c.replace('\n', ''))
        elif c.name == 'b' or c.name == 'strong':
            bold(p, c.text)
        elif c.name == 'italic' or c.name == 'em':
            italic(p, c.text)
        elif c.name == 'code':
            note_code(p, c.text)
    document.add_paragraph()


def caution(document, el):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cells([cell], "#feefe3")
    p = cell.paragraphs[0]
    p.style = CAUTION_STYLE
    pf = p.paragraph_format
    pf.space_before = Pt(15)
    pf.space_after = Pt(15)
    pf.left_indent = Pt(20)
    pf.right_indent = Pt(15)

    for c in el.contents:
        if isinstance(c, str):
            p.add_run(c.replace('\n', ' '))
        elif c.name == 'b' or c.name == 'strong':
            bold(p, c.text)
        elif c.name == 'italic' or c.name == 'em':
            italic(p, c.text)
        elif c.name == 'code':
            caution_code(p, c.text)
    document.add_paragraph()


def h1(document, text):
    p = document.add_paragraph(text)
    p.style = HEADING_1


def h2(document, text):
    p = document.add_paragraph(text)
    p.style = HEADING_2


def h3(document, text):
    p = document.add_paragraph(text)
    p.style = HEADING_3


def devsite_code(document, el):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cells([cell], "#f1f3f4")
    p = cell.paragraphs[0]
    p.style = NOTE_STYLE
    pf = p.paragraph_format
    pf.space_before = Pt(15)
    pf.space_after = Pt(15)
    pf.left_indent = Pt(20)
    pf.right_indent = Pt(15)

    text = el.text.replace('<br/>', '\n').replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>')
    p.add_run(text)
    # for span in el.find_all('span'):
    #     if span.has_attr('class'):
    #         if span['class'][0] == 'com':
    #             p.add_run(text).style = CODE_PINK
    #
    #         elif span['class'][0] == 'lit' or span['class'][0] == 'dec':
    #             p.add_run(text).style = CODE_MAROON
    #
    #         elif span['class'][0] == 'kwg' or span['class'][0] == 'tag':
    #             p.add_run(text).style = CODE_LIGHT_BLUE
    #
    #         elif span['class'][0] == 'str' or span['class'][0] == 'atv':
    #             p.add_run(text).style = CODE_DARK_GREEN
    #
    #         elif span['class'][0] == 'typ' or span['class'][0] == 'atn':
    #             p.add_run(text).style = CODE_PURPLE
    document.add_paragraph()


def bold(p, text):
    p.add_run(text.replace('\n', ' ') + ' ').bold = True


def italic(p, text):
    p.add_run(text.replace('\n', ' ') + ' ').italic = True


def caution_code(p, text):
    p.add_run(text.replace('\n', ' ') + ' ').style = CAUTION_CODE


def note_code(p, text):
    p.add_run(text.replace('\n', ' ') + ' ').style = NOTE_CODE


def code(p, text):
    p.add_run(text.replace('\n', ' ') + ' ').style = CODE_BLUE


def img(document, src):
    if src[0] == '/':
        src = 'https://developer.android.com' + src
    urlarray = src.split('/')
    filename = urlarray[len(urlarray)-1]
    if path.exists(filename):
        document.add_picture(filename)
    else:
        file = open(filename, 'wb')
        file.write(requests.get(src).content)
        file.close()
        # todo: get document width
        document.add_picture(filename, document.width())

